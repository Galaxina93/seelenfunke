import sys
import os
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

def set_cell_background(cell, fill_color):
    """Sets background shading color for a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Sets cell padding in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{margin}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def set_table_borders(table, color="E2E8F0", sz="4", val="single"):
    """Sets light gray borders for the entire table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_callout_left_border(cell, color="C5A059", sz="24"):
    """Sets a thick left border for callout blocks, removes other borders."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Thick left border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:color'), color)
    tcBorders.append(left)
    
    # Remove others
    for b in ['top', 'bottom', 'right']:
        none = OxmlElement(f'w:{b}')
        none.set(qn('w:val'), 'none')
        tcBorders.append(none)
        
    tcPr.append(tcBorders)

def add_page_number(run):
    """Adds a dynamic page number field to a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_formatted_text(paragraph, text, default_font_size=11, font_color=None, default_font_name='Calibri'):
    """Parses markdown-like inline styles (**bold**, *italic*, `code`) and adds runs."""
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.name = default_font_name
        run.font.size = Pt(default_font_size)
        if font_color:
            run.font.color.rgb = font_color
            
        if part.startswith('***') and part.endswith('***'):
            run.text = part[3:-3]
            run.font.bold = True
            run.font.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(default_font_size - 1.5)
            # Slight gray color to indicate code
            run.font.color.rgb = RGBColor(90, 90, 90)
        else:
            run.text = part

def parse_markdown_blocks(markdown_text):
    """Splits markdown into blocks: ('heading', level, text), ('list', type, text), ('table', rows), ('code', text), ('paragraph', text)"""
    lines = markdown_text.split('\n')
    blocks = []
    current_block = []
    in_code_block = False
    in_table = False

    for line in lines:
        cleaned_line = line.strip()
        
        # Code block toggle
        if line.startswith('```'):
            if in_code_block:
                blocks.append(('code', '\n'.join(current_block)))
                current_block = []
                in_code_block = False
            else:
                if current_block:
                    blocks.append(('paragraph', ' '.join(current_block)))
                current_block = []
                in_code_block = True
            continue
            
        if in_code_block:
            current_block.append(line)
            continue
            
        # Table detection
        if line.startswith('|'):
            if not in_table:
                if current_block:
                    blocks.append(('paragraph', ' '.join(current_block)))
                current_block = [line]
                in_table = True
            else:
                current_block.append(line)
            continue
        elif in_table:
            # End of table
            blocks.append(('table', parse_markdown_table(current_block)))
            current_block = []
            in_table = False
            
        # Empty line ends current paragraph
        if cleaned_line == '':
            if current_block:
                blocks.append(('paragraph', ' '.join(current_block)))
                current_block = []
            continue
            
        # Headings
        if cleaned_line.startswith('#'):
            if current_block:
                blocks.append(('paragraph', ' '.join(current_block)))
                current_block = []
            level = len(cleaned_line) - len(cleaned_line.lstrip('#'))
            blocks.append(('heading', level, cleaned_line.lstrip('#').strip()))
            continue
            
        # Lists
        if cleaned_line.startswith('- ') or cleaned_line.startswith('* '):
            if current_block:
                blocks.append(('paragraph', ' '.join(current_block)))
                current_block = []
            blocks.append(('list', 'bullet', cleaned_line[2:].strip()))
            continue
            
        if re.match(r'^\d+\.\s', cleaned_line):
            if current_block:
                blocks.append(('paragraph', ' '.join(current_block)))
                current_block = []
            match = re.match(r'^(\d+)\.\s(.*)', cleaned_line)
            blocks.append(('list', 'numbered', match.group(2).strip()))
            continue

        # Standard paragraph accumulator
        current_block.append(line)

    if current_block:
        if in_table:
            blocks.append(('table', parse_markdown_table(current_block)))
        elif in_code_block:
            blocks.append(('code', '\n'.join(current_block)))
        else:
            blocks.append(('paragraph', ' '.join(current_block)))

    return blocks

def parse_markdown_table(lines):
    """Parses markdown table lines into a list of list of cells, stripping headers separator."""
    rows = []
    for line in lines:
        if not line.strip():
            continue
        # Split by | and filter outer items
        cells = [c.strip() for c in line.split('|')]
        # If line starts and ends with |, cells[0] and cells[-1] will be empty
        if line.strip().startswith('|'):
            cells = cells[1:]
        if line.strip().endswith('|'):
            cells = cells[:-1]
            
        # Skip separator row like |---|---|
        if cells and all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
            
        rows.append(cells)
    return rows

def generate_document(data):
    title = data.get('title', 'Bericht')
    content_markdown = data.get('content_markdown', '')
    design = data.get('design', 'seelenfunke')
    agent_name = data.get('agentName', 'System')
    logo_path = data.get('logo_path', None)
    output_path = data.get('output_path', 'report.docx')
    owner_name = data.get('owner_name', 'Mein Seelenfunke')
    
    footer_info = data.get('footer_info', {})
    company_name = footer_info.get('company_name', owner_name)
    owner_proprietor = footer_info.get('owner_proprietor', 'Alina Steinhauer')
    company_street = footer_info.get('company_street', 'Carl-Goerdeler-Ring')
    company_street_number = footer_info.get('company_street_number', '26')
    company_zip = footer_info.get('company_zip', '38518')
    company_city = footer_info.get('company_city', 'Gifhorn')
    company_email = footer_info.get('company_email', 'kontakt@mein-seelenfunke.de')
    owner_website = footer_info.get('owner_website', 'www.mein-seelenfunke.de')
    owner_iban = footer_info.get('owner_iban', 'Wird nachgereicht')
    owner_tax_id = footer_info.get('owner_tax_id', '')
    owner_ust_id = footer_info.get('owner_ust_id', '')
    owner_court = footer_info.get('owner_court', 'Gifhorn')

    doc = Document()

    # Layout Setup
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Style Configurations
    # Define color scheme
    if design == 'seelenfunke':
        color_primary = RGBColor(197, 160, 89)   # Gold #C5A059
        color_secondary = RGBColor(17, 24, 39)   # Dark Gray/Black #111827
        color_text = RGBColor(45, 55, 72)        # Slate #2D3748
        bg_header_hex = "111827"                 # Dark Gray/Black for tables
        bg_zebra_hex = "FBF9F4"                  # Soft gold-gray zebra tint
    else:
        color_primary = RGBColor(31, 78, 120)    # Corporate Steel Blue #1F4E78
        color_secondary = RGBColor(43, 108, 176) # Secondary Blue #2B6CB0
        color_text = RGBColor(45, 55, 72)        # Dark Slate #2D3748
        bg_header_hex = "1F4E78"
        bg_zebra_hex = "F7FAFC"                  # Neutral zebra tint

    # Set normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = color_text

    # 1. Header & Footer Design with First Page Difference
    first_section = doc.sections[0]
    
    if design == 'seelenfunke':
        # Enable different headers/footers for first page
        first_section.different_first_page_header_footer = True
        
        # Header for subsequent pages (page 2+)
        header = first_section.header
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.0))
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Make header table borderless
        tblPr = header_table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Left cell: Title & Company
        left_cell = header_table.rows[0].cells[0]
        left_p = left_cell.paragraphs[0]
        left_run = left_p.add_run(title)
        left_run.font.name = 'Georgia'
        left_run.font.bold = True
        left_run.font.size = Pt(9.0)
        left_run.font.color.rgb = color_primary
        
        meta_run = left_cell.add_paragraph().add_run(company_name)
        meta_run.font.name = 'Calibri'
        meta_run.font.size = Pt(8.0)
        meta_run.font.color.rgb = RGBColor(120, 120, 120)

        # Right cell: Logo PNG (slightly larger for subsequent page headers)
        right_cell = header_table.rows[0].cells[1]
        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if logo_path and os.path.exists(logo_path):
            try:
                right_p.add_run().add_picture(logo_path, width=Cm(3.2))
            except Exception:
                right_p.text = company_name
                right_p.runs[0].font.name = 'Georgia'
                right_p.runs[0].font.bold = True
                right_p.runs[0].font.color.rgb = color_primary
        else:
            right_p.text = company_name
            right_p.runs[0].font.name = 'Georgia'
            right_p.runs[0].font.bold = True
            right_p.runs[0].font.color.rgb = color_primary

        # Footer for subsequent pages (page 2+) - Full Corporate Footer
        footer = first_section.footer
        footer_table = footer.add_table(rows=2, cols=3, width=Inches(6.0))
        
        # Make footer table borderless except a top separator line
        tblPr = footer_table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'single')
        top_border.set(qn('w:sz'), '4') # thin line
        top_border.set(qn('w:space'), '0')
        top_border.set(qn('w:color'), 'E2E8F0') # Light gray line
        tblBorders.append(top_border)
        
        for b in ['left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Left column (Company Address)
        c_left = footer_table.rows[0].cells[0]
        p_left = c_left.paragraphs[0]
        p_left.paragraph_format.space_after = Pt(0)
        p_left.paragraph_format.line_spacing = 1.0
        
        r_addr = p_left.add_run(f"{company_name}\nInh. {owner_proprietor}\n{company_street} {company_street_number}\n{company_zip} {company_city}")
        r_addr.font.name = 'Calibri'
        r_addr.font.size = Pt(8.0)
        r_addr.font.color.rgb = RGBColor(120, 120, 120)
        
        # Center column (Contact Info)
        c_center = footer_table.rows[0].cells[1]
        p_center = c_center.paragraphs[0]
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_center.paragraph_format.space_after = Pt(0)
        p_center.paragraph_format.line_spacing = 1.0
        
        r_contact = p_center.add_run(f"E-Mail:\n{company_email}\n\nWeb:\n{owner_website}")
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(8.0)
        r_contact.font.color.rgb = RGBColor(120, 120, 120)
        
        # Right column (Tax & Bank Info)
        c_right = footer_table.rows[0].cells[2]
        p_right = c_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.space_after = Pt(0)
        p_right.paragraph_format.line_spacing = 1.0
        
        tax_str = f"St.-Nr.: {owner_tax_id}"
        if owner_ust_id:
            tax_str += f"\nUSt-IdNr.: {owner_ust_id}"
        tax_str += f"\nGerichtsstand: {owner_court}"
        
        r_bank = p_right.add_run(f"IBAN:\n{owner_iban}\n\n{tax_str}")
        r_bank.font.name = 'Calibri'
        r_bank.font.size = Pt(8.0)
        r_bank.font.color.rgb = RGBColor(120, 120, 120)
        
        # Page number row (Row 1, Center cell)
        c_page = footer_table.rows[1].cells[1]
        p_page = c_page.paragraphs[0]
        p_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_page.paragraph_format.space_before = Pt(4)
        p_page.paragraph_format.space_after = Pt(0)
        
        run_page = p_page.add_run("Seite ")
        run_page.font.name = 'Calibri'
        run_page.font.size = Pt(8.0)
        run_page.font.color.rgb = RGBColor(120, 120, 120)
        add_page_number(run_page)
    else:
        # Generic Header (No Logo, No text)
        first_section.different_first_page_header_footer = True
        
        # Subsequent page footer page number
        footer = first_section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_page = footer_p.add_run("Seite ")
        run_page.font.name = 'Calibri'
        run_page.font.size = Pt(8.0)
        run_page.font.color.rgb = RGBColor(120, 120, 120)
        add_page_number(run_page)

    # 3. Content Body
    # Large Logo at the top of the first page body for Seelenfunke
    if design == 'seelenfunke':
        if logo_path and os.path.exists(logo_path):
            logo_p = doc.add_paragraph()
            logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_p.paragraph_format.space_before = Pt(0)
            logo_p.paragraph_format.space_after = Pt(24)
            try:
                logo_p.add_run().add_picture(logo_path, width=Cm(5.5))
            except Exception:
                pass

    # Main Document H1 Title (Cover-style Heading)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run(title)
    title_run.font.name = 'Georgia' if design == 'seelenfunke' else 'Calibri Light'
    title_run.font.size = Pt(24) if design == 'seelenfunke' else Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = color_secondary
    title_run.font.bold = True
    title_run.font.color.rgb = color_secondary
    
    # Underline rule below title
    if design == 'seelenfunke':
        # Add H1 gold accent bar below title
        rule_p = doc.add_paragraph()
        rule_p.paragraph_format.space_after = Pt(18)
        # We can implement a clean thin table row colored in Gold as a rule
        rule_table = doc.add_table(rows=1, cols=1)
        rule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = rule_table.rows[0].cells[0]
        cell.width = Inches(6.0)
        set_cell_background(cell, "C5A059")
        set_cell_margins(cell, top=20, bottom=20, left=0, right=0)
        # Clear borders
        tblPr = rule_table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    # Parse Markdown content
    blocks = parse_markdown_blocks(content_markdown)

    for block in blocks:
        block_type = block[0]
        
        if block_type == 'heading':
            level = block[1]
            text = block[2]
            
            # Map levels nicely
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            if level == 1:
                p.paragraph_format.space_before = Pt(20)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(text)
                run.font.name = 'Georgia' if design == 'seelenfunke' else 'Calibri Light'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = color_primary
                
                # Bottom border line
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '4')
                bottom.set(qn('w:color'), bg_header_hex)
                pBdr.append(bottom)
                pPr.append(pBdr)
                
            elif level == 2:
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                run.font.name = 'Georgia' if design == 'seelenfunke' else 'Calibri Light'
                run.font.size = Pt(13.5)
                run.font.bold = True
                run.font.color.rgb = color_primary
            else:
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(text)
                run.font.name = 'Georgia' if design == 'seelenfunke' else 'Calibri'
                run.font.size = Pt(11.5)
                run.font.bold = True
                run.font.color.rgb = color_secondary
                
        elif block_type == 'list':
            list_type = block[1]
            text = block[2]
            
            p = doc.add_paragraph(style='List Bullet' if list_type == 'bullet' else 'List Number')
            p.paragraph_format.space_after = Pt(3)
            add_formatted_text(p, text, default_font_size=10.5, font_color=color_text, default_font_name='Calibri')
            
        elif block_type == 'table':
            table_data = block[1]
            if not table_data:
                continue
                
            num_cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=0, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            
            border_color = "E2E8F0"
            if design == 'seelenfunke':
                border_color = "E5D3B3" # Elegant light gold/champagne table borders
            set_table_borders(table, color=border_color)
            
            # Populate table rows
            for row_idx, row in enumerate(table_data):
                row_cells = table.add_row().cells
                is_header = (row_idx == 0)
                
                # Fill cells
                for col_idx, cell_text in enumerate(row):
                    if col_idx >= num_cols:
                        break
                    cell = row_cells[col_idx]
                    set_cell_margins(cell)
                    
                    # Align numbers to the right, headers center, others left
                    align = WD_ALIGN_PARAGRAPH.LEFT
                    if is_header:
                        align = WD_ALIGN_PARAGRAPH.CENTER
                    elif re.match(r'^-?[\d\.,\s]+[€%]?$', cell_text.strip()):
                        align = WD_ALIGN_PARAGRAPH.RIGHT
                        
                    # Color styling
                    if is_header:
                        set_cell_background(cell, bg_header_hex)
                        fill_color = RGBColor(255, 255, 255)
                    else:
                        fill_color = color_text
                        if row_idx % 2 == 0:
                            set_cell_background(cell, bg_zebra_hex)
                            
                    p_cell = cell.paragraphs[0]
                    p_cell.alignment = align
                    add_formatted_text(p_cell, cell_text, default_font_size=9.5, font_color=fill_color, default_font_name='Calibri')
                    if is_header:
                        for r in p_cell.runs:
                            if design == 'seelenfunke':
                                r.font.name = 'Georgia'
                            r.font.bold = True
            
            # Spacer after table
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)
            spacer.paragraph_format.line_spacing = Pt(2)
            run = spacer.add_run()
            run.font.size = Pt(2)
            
        elif block_type == 'code':
            code_text = block[1]
            
            # Styled callout/code container using a single cell table
            code_table = doc.add_table(rows=1, cols=1)
            code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = code_table.rows[0].cells[0]
            cell.width = Inches(6.0)
            
            set_cell_background(cell, "F7FAFC")
            set_cell_margins(cell, top=180, bottom=180, left=220, right=220)
            set_callout_left_border(cell, color="4A5568", sz="12") # Steel border
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            
            # Write lines
            lines = code_text.split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    p_cell = cell.add_paragraph()
                    p_cell.paragraph_format.space_after = Pt(0)
                run = p_cell.add_run(line)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(74, 85, 104)
                
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
        elif block_type == 'paragraph':
            text = block[1]
            
            # Check if block is a callout quote (starts with "> ")
            if text.startswith('&gt;') or text.startswith('>'):
                # Callout block
                cleaned_text = text.lstrip('&gt;').lstrip('>').strip()
                
                callout_table = doc.add_table(rows=1, cols=1)
                callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = callout_table.rows[0].cells[0]
                cell.width = Inches(6.0)
                
                bg_callout = "FFFDF5" if design == 'seelenfunke' else "F7FAFC"
                border_color = "C5A059" if design == 'seelenfunke' else "2B6CB0"
                
                set_cell_background(cell, bg_callout)
                set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
                set_callout_left_border(cell, color=border_color, sz="24")
                
                p_cell = cell.paragraphs[0]
                p_cell.paragraph_format.space_after = Pt(0)
                
                # Premium styling for blockquote text: Georgia font, slightly smaller, italicized
                if design == 'seelenfunke':
                    add_formatted_text(p_cell, cleaned_text, default_font_size=10.0, font_color=RGBColor(80, 70, 50), default_font_name='Georgia')
                    for run in p_cell.runs:
                        run.font.italic = True
                else:
                    add_formatted_text(p_cell, cleaned_text, default_font_size=10.0, font_color=color_text, default_font_name='Calibri')
                    for run in p_cell.runs:
                        run.font.italic = True
                
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                add_formatted_text(p, text, default_font_size=10.5, font_color=color_text, default_font_name='Calibri')

    # Save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Success: Document generated at {output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Error: Missing JSON arguments file path.")
        sys.exit(1)
        
    args_file = sys.argv[1]
    if not os.path.exists(args_file):
        print(f"Error: Arguments file {args_file} does not exist.")
        sys.exit(1)
        
    with open(args_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    try:
        generate_document(data)
    except Exception as e:
        import traceback
        traceback.print_exc()
        sys.exit(1)
